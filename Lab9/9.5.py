from Lab8.catalog import read_data, media_gen
from docx import Document


def create_docx(cat: list, sab: list):
    for x in range(len(cat)):
        doc = Document()
        doc.add_heading('DIPLOMA', 0)
        p = doc.add_paragraph(f'{sab[0]}  ')
        p.add_run(cat[x][0]).bold = True
        p = doc.add_paragraph(f'{sab[1]}  ')
        p.add_run(cat[x][1]).italic = True
        doc.add_paragraph(f'{sab[2]}  {cat[x][2]}')
        doc.save(f'premiul-{x + 1}.docx')


with open('diploma.txt', 'r') as dpl:
    sablon = [x.rstrip('\n{}') for x in dpl.readlines()]


with open('diploma.txt', 'r') as dpl:
    sablon1 = dpl.read()


catalog = media_gen(read_data('..//Lab7/note.txt'), read_data('..//Lab7/teze.txt'))

ordered_cat = sorted(catalog.items(), key=lambda v: v[1], reverse=True)

premianti = [[(x + 1) * 'I', ordered_cat[x][0], f'{ordered_cat[x][1]:.2f}'] for x in range(3)]

create_docx(premianti, sablon)


for i, e in enumerate(ordered_cat, start=1):
    if i < 4:
        with open(f'premiul-{i}.txt', 'w') as file:
            print(sablon1.format((i * 'I'), e[0], f'{e[1]:.2f}'), file=file)
